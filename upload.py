import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
import docx
from docx import Document
import io
from PIL import Image
import tempfile
import os
import tabula
import json
import cv2
import numpy as np
import librosa

def clear_temp_files():
    temp_dir = tempfile.gettempdir()
    for filename in os.listdir(temp_dir):
        if filename.startswith('temp') and filename.endswith('.csv'):
            os.remove(os.path.join(temp_dir, filename))

clear_temp_files()

def preprocess_csv(df):
    df = df.convert_dtypes()
    return df

def read_pdf(file):
    file.seek(0)
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    return text

def extract_tables_from_pdf(file):
    file.seek(0)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
        temp_pdf.write(file.read())
        temp_pdf.flush()
        encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
        for encoding in encodings:
            try:
                tables = tabula.read_pdf(temp_pdf.name, pages='all', multiple_tables=True, encoding=encoding)
                return tables
            except Exception as e:
                continue
        st.error(f"Error extracting tables from PDF: Unable to decode with available encodings")
        return []

def extract_images_from_pdf(file):
    file.seek(0)
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    images = []
    
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        image_list = page.get_images(full=True)
        
        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))
            images.append(image)
    
    return images

def read_docx(file):
    file.seek(0)
    doc = docx.Document(file)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

def extract_tables_from_docx(file):
    file.seek(0)
    doc = docx.Document(file)
    tables = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        df = pd.DataFrame(data)
        tables.append(df)
    return tables

def extract_images_from_docx(file):
    file.seek(0)
    doc = Document(file)
    images = []
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_part = rel.target_part
            image_bytes = image_part.blob
            image = Image.open(io.BytesIO(image_bytes))
            images.append(image)
    
    return images

def read_txt(file):
    file.seek(0)
    return file.read().decode("utf-8")

def display_content_with_tables(text, tables):
    st.markdown(text)
    for idx, table in enumerate(tables):
        st.write(f"**Table {idx+1}:**")
        st.write(table)

def preprocess_and_update_pdf(file):
    tables = extract_tables_from_pdf(file)
    preprocessed_tables = []
    
    for df in tables:
        if not df.empty:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as temp_file:
                df.to_csv(temp_file.name, index=False)
                temp_df = pd.read_csv(temp_file.name)
                preprocessed_df = preprocess_csv(temp_df)
                preprocessed_tables.append(preprocessed_df)
    
    return preprocessed_tables

def preprocess_and_update_docx(file):
    tables = extract_tables_from_docx(file)
    preprocessed_tables = []

    for df in tables:
        if not df.empty:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as temp_file:
                df.to_csv(temp_file.name, index=False)
                temp_df = pd.read_csv(temp_file.name)
                preprocessed_df = preprocess_csv(temp_df)
                preprocessed_tables.append(preprocessed_df)
    
    return preprocessed_tables

def handle_file_upload(page_name):
    # Generate a unique key based on the page name
    key = f"file_uploader_{page_name}"

    if 'uploaded_files' not in st.session_state:
        st.session_state['uploaded_files'] = None

    uploaded_files = st.file_uploader("Upload files", type=["csv", "xlsx", "json", "txt", "pdf", "docx", "jpg", "jpeg", "png", "mp4", "mp3", "wav"], accept_multiple_files=True, key=key)

    if uploaded_files:
        st.session_state['uploaded_files'] = uploaded_files
        dfs_tabular = []
        dfs_non_tabular = []
        
        for uploaded_file in uploaded_files:
            file_extension = uploaded_file.name.split(".")[-1].lower()

            try:
                if file_extension == 'csv':
                    encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
                    for encoding in encodings:
                        try:
                            df = pd.read_csv(uploaded_file, encoding=encoding)
                            if df.empty:
                                st.warning(f"The CSV file '{uploaded_file.name}' is empty.")
                                break
                            dfs_tabular.append(preprocess_csv(df))
                            break
                        except pd.errors.EmptyDataError:
                            st.error(f"No columns to parse from file '{uploaded_file.name}'.")
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        st.error(f"Unable to read the file with any of the attempted encodings: {encodings}")
                        return
                elif file_extension == 'xlsx':
                    df = pd.read_excel(uploaded_file)
                    dfs_tabular.append(preprocess_csv(df))
                elif file_extension == 'json':
                    df = pd.read_json(uploaded_file)
                    dfs_tabular.append(preprocess_csv(df))
                elif file_extension == 'pdf':
                    tables = extract_tables_from_pdf(uploaded_file)
                    for table in tables:
                        dfs_tabular.append(preprocess_csv(table))
                    text = read_pdf(uploaded_file)
                    dfs_non_tabular.append(text)
                elif file_extension == 'docx':
                    tables = extract_tables_from_docx(uploaded_file)
                    for table in tables:
                        dfs_tabular.append(preprocess_csv(table))
                    text = read_docx(uploaded_file)
                    dfs_non_tabular.append(text)
                elif file_extension == 'txt':
                    text = read_txt(uploaded_file)
                    dfs_non_tabular.append(text)
                elif file_extension in ['jpg', 'jpeg', 'png']:
                    image = Image.open(uploaded_file)
                    dfs_non_tabular.append(image)
                elif file_extension in ['mp4']:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp_file:
                        temp_file.write(uploaded_file.read())
                        temp_file.flush()
                        video = cv2.VideoCapture(temp_file.name)
                        # Extract video properties as needed
                        duration = video.get(cv2.CAP_PROP_FRAME_COUNT) / video.get(cv2.CAP_PROP_FPS)
                        summary = f"Video duration: {duration:.2f} seconds"
                        dfs_non_tabular.append(summary)
                elif file_extension in ['mp3', 'wav']:
                    y, sr = librosa.load(uploaded_file)
                    duration = librosa.get_duration(y=y, sr=sr)
                    summary = f"Audio duration: {duration:.2f} seconds"
                    dfs_non_tabular.append(summary)
                else:
                    st.warning(f"Unsupported file type: {uploaded_file.name}")

            except Exception as e:
                st.error(f"An error occurred while processing the file '{uploaded_file.name}': {str(e)}")
                continue
            
        st.session_state['dfs_tabular'] = dfs_tabular
        st.session_state['dfs_non_tabular'] = dfs_non_tabular

    return st.session_state['uploaded_files'], st.session_state['dfs_tabular'], st.session_state['dfs_non_tabular']
