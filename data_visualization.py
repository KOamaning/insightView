import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
import docx
from docx import Document
import io
from PIL import Image
import tempfile
from PyPDF2 import PdfReader
from sklearn.linear_model import LinearRegression
import tabula
import base64
import os
from pygwalker.api.streamlit import StreamlitRenderer
from pandasai import SmartDataframe
from pandasai import Agent
import numpy as np
from scipy import stats      
from io import BytesIO
from sklearn.impute import SimpleImputer
from sklearn.experimental import enable_iterative_imputer
from sklearn.impute import IterativeImputer
from sklearn.linear_model import BayesianRidge
from upload import handle_file_upload
from dotenv import load_dotenv

PANDASAI_API_KEY = os.getenv('PANDASAI_API_KEY')

# Initialize session state variables
if 'processed_df' not in st.session_state:
                st.session_state.processed_df = None
if 'data_processed' not in st.session_state:
                st.session_state.data_processed = False
if 'uploaded_files' not in st.session_state:
                st.session_state.uploaded_files = None

def configure():
        load_dotenv()

            # Clear temporary files at the start
def clear_temp_files():
                temp_dir = tempfile.gettempdir()
                for filename in os.listdir(temp_dir):
                    if filename.startswith('temp') and filename.endswith('.csv'):
                        os.remove(os.path.join(temp_dir, filename))

clear_temp_files()

def preprocess_csv(df):
                # Add your preprocessing steps here
                # For example, converting data types, handling missing values, etc.
                df = df.convert_dtypes()
                return df

def read_pdf(file):
                file.seek(0)
                pdf_document = fitz.open(stream=file.read(), filetype="pdf")
                text = ""
                for page_num in range(pdf_document.page_count):
                    page = pdf_document.load_page(page_num)
                    text += page.get_text()
                return text


def extract_tables_from_pdf(file):
                file.seek(0)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                    temp_pdf.write(file.read())
                    temp_pdf.flush()
                    encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
                    for encoding in encodings:
                        try:
                            tables = tabula.read_pdf(temp_pdf.name, pages='all', multiple_tables=True, encoding=encoding)
                            return tables
                        except Exception as e:
                            continue
                    st.error(f"Error extracting tables from PDF: Unable to decode with available encodings")
                    return []


def extract_images_from_pdf(file):
                file.seek(0)
                pdf_document = fitz.open(stream=file.read(), filetype="pdf")
                images = []
                
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    image_list = page.get_images(full=True)
                    
                    for img_index, img_info in enumerate(image_list):
                        xref = img_info[0]
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        image = Image.open(io.BytesIO(image_bytes))
                        images.append(image)
                
                return images


def read_docx(file):
                file.seek(0)
                doc = docx.Document(file)
                text = []
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                return "\n".join(text)

def extract_tables_from_docx(file):
                file.seek(0)
                doc = docx.Document(file)
                tables = []
                for table in doc.tables:
                    data = []
                    for row in table.rows:
                        data.append([cell.text for cell in row.cells])
                    df = pd.DataFrame(data)
                    tables.append(df)
                return tables


def extract_images_from_docx(file):
                file.seek(0)
                doc = Document(file)
                images = []
                
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        image = Image.open(io.BytesIO(image_bytes))
                        images.append(image)
                
                return images


def read_txt(file):
                file.seek(0)
                return file.read().decode("utf-8")

def display_content_with_tables(text, tables):
                st.markdown(text)
                for idx, table in enumerate(tables):
                    st.write(f"**Table {idx+1}:**")
                    st.write(table)

def preprocess_and_update_pdf(file):
                tables = extract_tables_from_pdf(file)
                preprocessed_tables = []
                
                for df in tables:
                    if not df.empty:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as temp_file:
                            df.to_csv(temp_file.name, index=False)
                            temp_df = pd.read_csv(temp_file.name)
                            preprocessed_df = preprocess_csv(temp_df)
                            preprocessed_tables.append(preprocessed_df)
                
                return preprocessed_tables

def preprocess_and_update_docx(file):
                tables = extract_tables_from_docx(file)
                preprocessed_tables = []

                for df in tables:
                    if not df.empty:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as temp_file:
                            df.to_csv(temp_file.name, index=False)
                            temp_df = pd.read_csv(temp_file.name)
                            preprocessed_df = preprocess_csv(temp_df)
                            preprocessed_tables.append(preprocessed_df)
                
                return preprocessed_tables


def data_visualization():
    configure()
    st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Data Visualization</h1>', unsafe_allow_html=True)
    uploaded_files, dfs_tabular, dfs_non_tabular = handle_file_upload(page_name=data_visualization)

    if dfs_tabular:
        # Combine all DataFrames in dfs_tabular
        final_df = pd.concat(dfs_tabular, ignore_index=True)

        # Use a temporary file instead of a hard-coded path
        with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as temp_file:
            output_path = temp_file.name
            final_df.to_csv(output_path, index=False)

        # Read the data back from the CSV file into the df variable
        df = pd.read_csv(output_path)

        try:
            sdf = SmartDataframe(output_path)
            st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Graph generation using AI</h1>', unsafe_allow_html=True)
            query = st.text_input("Enter prompt to generate graph. eg: Plot a bar graph for column 1", "")
            if st.button("Generate graph", type="primary"):
                if query:
                    response = sdf.chat(query)
                    image = response
                    st.image(image)
                else:
                    st.warning("Please enter a query before submitting.")
        except Exception as e:
            st.error(f"An error occurred while generating the graph: {str(e)}")

        

        st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Visualization center</h1>', unsafe_allow_html=True)
        st.markdown('<p style="font-size: 0.6rem; margin-top: 0;">Data Tab: provides general overview of your data. Visualization Tab: provides a wide array of data visualization options using drag and drop</p>', unsafe_allow_html=True)

        try:
           
            def get_pyg_renderer() -> "StreamlitRenderer":
                return StreamlitRenderer(df, spec="./gw_config.json")
            
            renderer = get_pyg_renderer()      
            renderer.explorer(default_tab="data")
        except Exception as e:
            st.error(f"An error occurred while rendering the visualization: {str(e)}")

        # Clean up the temporary file
        os.unlink(output_path)
    else:
        st.warning("Please upload a file that contains tabular data for data visualization.")