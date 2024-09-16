import streamlit as st
import google.generativeai as genai
from pypdf import PdfReader 
import os
from fpdf import FPDF
import pandas as pd
import docx
from docx import Document
import tempfile
import fitz
from transformers import pipeline,AutoTokenizer, AutoModelForSeq2SeqLM
from textwrap import wrap
from docx import Document
import io
from upload import handle_file_upload

# Initialize session state variables
if 'processed_df' not in st.session_state:
                st.session_state.processed_df = None
if 'data_processed' not in st.session_state:
                st.session_state.data_processed = False
if 'uploaded_files' not in st.session_state:
                st.session_state.uploaded_files = None
if 'dfs_tabular' not in st.session_state:
                st.session_state.dfs_tabular = None
if 'dfs_non_tabular' not in st.session_state:
                st.session_state.dfs_non_tabular = None
 # Clear temporary files at the start
def clear_temp_files():
                temp_dir = tempfile.gettempdir()
                for filename in os.listdir(temp_dir):
                    if filename.startswith('temp') and filename.endswith('.csv'):
                        os.remove(os.path.join(temp_dir, filename))

clear_temp_files()


def read_pdf(file):
                file.seek(0)
                pdf_document = fitz.open(stream=file.read(), filetype="pdf")
                text = ""
                for page_num in range(pdf_document.page_count):
                    page = pdf_document.load_page(page_num)
                    text += page.get_text()
                return text


def read_docx(file):
                file.seek(0)
                doc = docx.Document(file)
                text = []
                for paragraph in doc.paragraphs:
                    text.append(paragraph.text)
                return "\n".join(text)



def read_txt(file):
    file.seek(0)
    return file.read().decode("utf-8")


def get_or_download_model(model_name, local_dir):
    if not os.path.exists(local_dir):
        print(f"Model not found locally. Downloading {model_name}...")
        model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        
        os.makedirs(local_dir, exist_ok=True)
        model.save_pretrained(local_dir)
        tokenizer.save_pretrained(local_dir)
        print(f"Model and tokenizer saved to {local_dir}")
    else:
        print(f"Loading model from {local_dir}")
        model = AutoModelForSeq2SeqLM.from_pretrained(local_dir)
        tokenizer = AutoTokenizer.from_pretrained(local_dir)
    
    return model, tokenizer



# Usage in your main code
model_name = "facebook/bart-large-cnn"
local_dir = "./my_summarization_model"
model, tokenizer = get_or_download_model(model_name, local_dir)




def text_summarization():
    st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Text data summarization</h1>', unsafe_allow_html=True)
    uploaded_files, _, dfs_non_tabular = handle_file_upload(page_name=text_summarization)
    
    dfs_non_tabular = []
    
    if uploaded_files is None or len(uploaded_files) == 0:
        st.warning("Please upload a file that contains text data for text summarization.")
        return
    
    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name
        file_extension = file_name.split(".")[-1].lower()

        if file_extension not in ["txt", "pdf", "docx"]:
            st.warning(f"Unsupported file type for file {file_name}. Please upload a file that contains text data for text summarization.")
        else:
            if file_extension == 'pdf':
                text = read_pdf(uploaded_file)
            elif file_extension == 'docx':
                text = read_docx(uploaded_file)
            elif file_extension == 'txt':
                text = read_txt(uploaded_file)
            dfs_non_tabular.append(text)

    if dfs_non_tabular:
        combined_text = "\n\n--- New Document ---\n\n".join(dfs_non_tabular)

        # Initialize the summarization pipeline
        summarizer = pipeline("summarization", model=model, tokenizer=tokenizer)
        # Summarize the text, handling large inputs by chunking
        chunk_size = 1024
        chunks = wrap(combined_text, chunk_size)

        summaries = []
        for chunk in chunks:
            summary = summarizer(chunk, max_length=150, min_length=30, do_sample=False)
            summaries.append(summary[0]["summary_text"])

        final_summary = " ".join(summaries)

        st.write("Summary:")
        st.write(final_summary)

        # Initialize session state
        if 'show_download_options' not in st.session_state:
            st.session_state.show_download_options = False

        # Main download button
        if st.button("Download", key="main_download_button",type="primary"):
            st.session_state.show_download_options = True

        # Show download options if the button has been clicked
        if st.session_state.show_download_options:
            st.markdown("""
                <style>
                .stSelectbox {
                    max-width: 200px;
                }
                </style>
                """, unsafe_allow_html=True)
            
            file_format = st.selectbox('Select file format to download:', ['docx', 'pdf'], key="file_format_select")

            if st.button('Download Summary', key="download_summary_button",type="primary"):
                if file_format == 'docx':
                    doc = Document()
                    doc.add_heading('Text Summary', level=1)
                    doc.add_paragraph(final_summary)
                    
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    
                    st.download_button(
                        label="Download DOCX",
                        type="primary",
                        data=docx_buffer,
                        file_name="text_summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_docx_button"
                    )
                elif file_format == 'pdf':
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0, 10, txt="Text Summary", ln=1, align='C')
                    pdf.multi_cell(0, 10, txt=final_summary)
                    
                    pdf_buffer = io.BytesIO()
                    pdf_buffer.write(pdf.output(dest='S').encode('latin-1'))
                    pdf_buffer.seek(0)
                    
                    st.download_button(
                        label="Download PDF",
                        type="primary",
                        data=pdf_buffer,
                        file_name="text_summary.pdf",
                        mime="application/pdf",
                        key="download_pdf_button"
                    )
    