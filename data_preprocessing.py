import streamlit as st
import pandas as pd
import numpy as np
from scipy import stats
import fitz  # PyMuPDF
import docx
from docx import Document
import io
from PIL import Image
import tempfile
from PyPDF2 import PdfReader
from sklearn.linear_model import LinearRegression
import os
import tabula
import base64
from io import BytesIO
from sklearn.impute import SimpleImputer
from sklearn.experimental import enable_iterative_imputer
from sklearn.impute import IterativeImputer
from sklearn.linear_model import BayesianRidge
from mitosheet.streamlit.v1 import spreadsheet
from collections import OrderedDict
import matplotlib.pyplot as plt
from fpdf import FPDF
from upload import handle_file_upload



def save_dataframes_to_specific_csv(new_dfs, file_path):
    if isinstance(new_dfs, OrderedDict):
        combined_df = pd.DataFrame()  # Initialize an empty DataFrame
        for key, df in new_dfs.items():
            if isinstance(df, pd.DataFrame):
                combined_df = pd.concat([combined_df, df], ignore_index=True)
            else:
                st.warning(f"The item with key '{key}' is not a DataFrame and was not saved.")
        
        # Write the combined DataFrame to the specified CSV file
        combined_df.to_csv(file_path, index=False)
        st.success("Preprocessing complete. Proceed to the other pages")
    else:
        st.error("The provided data is not an OrderedDict.")

def convert_df(df, file_format):
    if file_format == 'csv':
        return df.to_csv(index=False).encode('utf-8')
    elif file_format == 'excel':
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Sheet1')
        return output.getvalue()
    elif file_format == 'json':
        return df.to_json().encode('utf-8')

def read_file_content(file_path):
    file_ext = os.path.splitext(file_path)[1].lower()
    if file_ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
    elif file_ext == ".docx":
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs])
    else:
        content = ""
    return content



def write_file_content(file_path, content):
    def try_encodings(text, encodings):
        for encoding in encodings:
            try:
                return text.encode(encoding)
            except UnicodeEncodeError:
                continue
        raise UnicodeEncodeError(f"None of the encodings worked: {encodings}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == ".txt":
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
    elif file_ext == ".docx":
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(file_path)
    elif file_ext == ".pdf":
        encodings = ["latin-1", "utf-8", "ascii"]
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        
        for line in content.split("\n"):
            encoded_line = try_encodings(line, encodings).decode("latin-1", "replace")
            pdf.multi_cell(0, 10, encoded_line)
        
        pdf.output(file_path)




def data_preprocessing():

    st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Data Preprocessing</h1>', unsafe_allow_html=True)
    uploaded_files, dfs_tabular, dfs_non_tabular = handle_file_upload(page_name="data_preprocessing")

    if not uploaded_files:
        st.warning("Please upload files to process.")
        return

    # Process tabular data
    if dfs_tabular:
        # Combine all DataFrames in dfs_tabular
        final_df = pd.concat(dfs_tabular, ignore_index=True)

        # Write the combined DataFrame to a CSV file
        output_path = r"C:\Users\Kwaku\Desktop\project_final\final_df.csv"
        final_df.to_csv(output_path, index=False)

        st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Preprocess Tabular Data:</h1>', unsafe_allow_html=True)
        # Use the spreadsheet function to edit the dataframe
        new_dfs, code = spreadsheet(final_df)

        if st.button("Update Tabular Data", key="update_tabular", type="primary"):
            if isinstance(new_dfs, OrderedDict) and len(new_dfs) > 0:
                current_df = list(new_dfs.values())[0]
                current_df.to_csv(output_path, index=False)
                st.success("Dataframe updated and saved.")
            else:
                st.warning("No changes detected in the dataframe.")
        
     
        # Initialize session state
        if 'show_download_options' not in st.session_state:
            st.session_state.show_download_options = False

        # Main download button for tabular data
        if st.button("Download Tabular Data", key="download_tabular", type="primary"):
            st.session_state.show_download_options = True

        # Show download options if the button has been clicked
        if st.session_state.show_download_options:
            # Apply custom CSS to reduce the width of the selectbox
            st.markdown("""
                <style>
                .stSelectbox {
                    max-width: 200px;
                }
                </style>
                """, unsafe_allow_html=True)
            
            # File download options
            file_format = st.selectbox('Select file format to download:', ['csv', 'excel', 'json'])
            
            # Use the most recent version of the dataframe for download
            if isinstance(new_dfs, OrderedDict) and len(new_dfs) > 0:
                download_df = list(new_dfs.values())[0]
            else:
                download_df = final_df
            
            file_data = convert_df(download_df, file_format)
            
            if file_format == 'csv':
                st.download_button(label="Download CSV", type="primary", data=file_data, file_name="data.csv", mime='text/csv')
            elif file_format == 'excel':
                st.download_button(label="Download Excel", type="primary", data=file_data, file_name="data.xlsx", mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            elif file_format == 'json':
                st.download_button(label="Download JSON", type="primary", data=file_data, file_name="data.json", mime='application/json')

        # Process non-tabular data (unchanged)
    if dfs_non_tabular:
        # Separate strings and images
        text_items = []
        image_items = []

        for item in dfs_non_tabular:
            if isinstance(item, str):
                text_items.append(item)
            elif isinstance(item, Image.Image):  # Check for image objects
                image_items.append(item)
            else:
                text_items.append(str(item))  # Convert other non-string items to string if necessary

        if text_items:  # Only process if there are text items
            # Join only the text items
            combined_text = "\n\n--- New Document ---\n\n".join(text_items)

            # Write the combined text to a file
            output_path = r"C:\Users\Kwaku\Desktop\project_final\upload_txt.txt"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(combined_text)

            # Read the content of the file
            file_content = read_file_content(output_path)

            st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Preprocess Text Data:</h1>', unsafe_allow_html=True)
            # Text area for editing the content
            updated_content = st.text_area("Edit the content below:", file_content, height=300)

            # Update button to save the content (for non-tabular data)
            if st.button("Update Text Data", key="update_text", type="primary"):
                write_file_content(output_path, updated_content)
                st.success(f"File has been updated and saved to {output_path}")

            # Initialize session state
            if 'show_download_options' not in st.session_state:
                st.session_state.show_download_options = False

            # Main download button for non-tabular data
            if st.button("Download Text Data", key="download_text", type="primary"):
                st.session_state.show_download_options = True

            # Show download options if the button has been clicked
            if st.session_state.show_download_options:
                # Apply custom CSS to reduce the width of the selectbox
                st.markdown("""
                    <style>
                    .stSelectbox {
                        max-width: 200px;
                    }
                    </style>
                    """, unsafe_allow_html=True)

                # Download the file
                download_format = st.selectbox("Select the format to download:", ("txt", "docx", "pdf"))

                download_path = os.path.splitext(output_path)[0] + f".{download_format}"
                write_file_content(download_path, updated_content)

                with open(download_path, "rb") as f:
                    st.download_button(
                        label="Download",
                        type="primary",
                        data=f,
                        file_name=os.path.basename(download_path),
                        mime="application/octet-stream"
                    )

    if not dfs_tabular and not dfs_non_tabular:
        st.warning("No data to process. Please upload valid files.")