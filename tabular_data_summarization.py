import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
import docx
from docx import Document
import io
from PIL import Image
import tempfile
from PyPDF2 import PdfReader
from sklearn.linear_model import LinearRegression
import os
import tabula
from lida import Manager, TextGenerationConfig, llm
import base64
from io import BytesIO
import os
from pandasai import SmartDataframe
from pandasai import Agent   
from upload import handle_file_upload   
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx.shared import Inches
from reportlab.lib.utils import ImageReader
import time
from dotenv import load_dotenv


    
# Initialize session state variables
if 'processed_df' not in st.session_state:
    st.session_state.processed_df = None
if 'data_processed' not in st.session_state:
    st.session_state.data_processed = False
if 'uploaded_files' not in st.session_state:
    st.session_state.uploaded_files = None
if 'dfs_tabular' not in st.session_state:
    st.session_state.dfs_tabular = None
if 'dfs_non_tabular' not in st.session_state:
    st.session_state.dfs_non_tabular = None

def configure():
        load_dotenv()

# Clear temporary files at the start
def clear_temp_files():
            temp_dir = tempfile.gettempdir()
            for filename in os.listdir(temp_dir):
                if filename.startswith('temp') and filename.endswith('.csv'):
                    os.remove(os.path.join(temp_dir, filename))

clear_temp_files()

def preprocess_csv(df):
            # Add your preprocessing steps here
            # For example, converting data types, handling missing values, etc.
            df = df.convert_dtypes()
            return df

def read_pdf(file):
            file.seek(0)
            pdf_document = fitz.open(stream=file.read(), filetype="pdf")
            text = ""
            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)
                text += page.get_text()
            return text


def extract_tables_from_pdf(file):
            file.seek(0)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                temp_pdf.write(file.read())
                temp_pdf.flush()
                encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
                for encoding in encodings:
                    try:
                        tables = tabula.read_pdf(temp_pdf.name, pages='all', multiple_tables=True, encoding=encoding)
                        return tables
                    except Exception as e:
                        continue
                st.error(f"Error extracting tables from PDF: Unable to decode with available encodings")
                return []


def extract_images_from_pdf(file):
            file.seek(0)
            pdf_document = fitz.open(stream=file.read(), filetype="pdf")
            images = []
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                image_list = page.get_images(full=True)
                
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))
                    images.append(image)
            
            return images



def read_docx(file):
            file.seek(0)
            doc = docx.Document(file)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return "\n".join(text)

def extract_tables_from_docx(file):
            file.seek(0)
            doc = docx.Document(file)
            tables = []
            for table in doc.tables:
                data = []
                for row in table.rows:
                    data.append([cell.text for cell in row.cells])
                df = pd.DataFrame(data)
                tables.append(df)
            return tables

def extract_images_from_docx(file):
            file.seek(0)
            doc = Document(file)
            images = []
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    image = Image.open(io.BytesIO(image_bytes))
                    images.append(image)
            
            return images


def read_txt(file):
            file.seek(0)
            return file.read().decode("utf-8")

def display_content_with_tables(text, tables):
            st.markdown(text)
            for idx, table in enumerate(tables):
                st.write(f"**Table {idx+1}:**")
                st.write(table)

def preprocess_and_update_pdf(file):
            tables = extract_tables_from_pdf(file)
            preprocessed_tables = []
            
            for df in tables:
                if not df.empty:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as temp_file:
                        df.to_csv(temp_file.name, index=False)
                        temp_df = pd.read_csv(temp_file.name)
                        preprocessed_df = preprocess_csv(temp_df)
                        preprocessed_tables.append(preprocessed_df)
            
            return preprocessed_tables

def preprocess_and_update_docx(file):
            tables = extract_tables_from_docx(file)
            preprocessed_tables = []

            for df in tables:
                if not df.empty:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as temp_file:
                        df.to_csv(temp_file.name, index=False)
                        temp_df = pd.read_csv(temp_file.name)
                        preprocessed_df = preprocess_csv(temp_df)
                        preprocessed_tables.append(preprocessed_df)
            
            return preprocessed_tables

def convert_to_docx(content, images):
    doc = Document()
    for item in content:
        if isinstance(item, str):
            doc.add_paragraph(item)
        elif isinstance(item, Image.Image):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                item.save(temp_img.name, format="PNG")
                doc.add_picture(temp_img.name, width=Inches(6))
                temp_img.close()
                # Add a small delay before trying to delete the file
                time.sleep(0.1)
                try:
                    os.unlink(temp_img.name)
                except PermissionError:
                    print(f"Warning: Unable to delete temporary file {temp_img.name}. It will be deleted when the system restarts.")
        doc.add_paragraph()  # Add space after each item
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def convert_to_pdf(content, images):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y_position = height - 72  # Start 1 inch from the top

    for item in content:
        if isinstance(item, str):
            textobject = c.beginText()
            textobject.setTextOrigin(72, y_position)
            textobject.setFont("Helvetica", 12)
            for line in item.split('\n'):
                textobject.textLine(line)
                y_position -= 14  # Adjust position for next line
            c.drawText(textobject)
        elif isinstance(item, Image.Image):
            img_width, img_height = item.size
            aspect = img_height / float(img_width)
            img_width = width - 144  # 2 inches margin
            img_height = img_width * aspect
            if y_position - img_height < 72:  # Check if image fits on page
                c.showPage()
                y_position = height - 72
            y_position -= img_height
            c.drawImage(ImageReader(item), 72, y_position, width=img_width, height=img_height)
        y_position -= 20  # Add space after each item
        if y_position < 72:  # Check if we need a new page
            c.showPage()
            y_position = height - 72
    c.save()
    return buffer.getvalue()
def tabular_data_summarization():
    configure()
    st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Tabular Data Summarization</h1>', unsafe_allow_html=True)
    
    # Initialize session state variables
    if 'content' not in st.session_state:
        st.session_state.content = []
    if 'images' not in st.session_state:
        st.session_state.images = []
    
    uploaded_files, dfs_tabular, _ = handle_file_upload(page_name="tabular_data_summarization")

    if not dfs_tabular:
        st.warning("Please upload a file that contains tabular data for tabular data summarization.")
        return

    # Combine all dataframes
    combined_df = pd.concat(dfs_tabular, ignore_index=True)

    # Display a sample of the data
    st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Sample of uploaded data:</h1>', unsafe_allow_html=True)
    st.write(combined_df.head())

    # Button to trigger summarization
    if st.button("Run Data Summarization", key="run_summarization", type="primary"):
        # Save combined dataframe to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as temp_file:
            combined_df.to_csv(temp_file.name, index=False)
            temp_file_path = temp_file.name

        try:
            api_key = os.getenv('api_key')
            if not api_key:
                raise ValueError("API key not found in environment variables")
            
            lida = Manager(text_gen=llm("cohere", api_key=api_key))
            textgen_config = TextGenerationConfig(n=3, temperature=0.5, model="command-r-plus", use_cache=True)     
            
            st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Summarization of your data</h1>', unsafe_allow_html=True)
        
            summary = lida.summarize(temp_file_path, summary_method="default", textgen_config=textgen_config)
            goals = lida.goals(summary, n=1, textgen_config=textgen_config)

            st.session_state.content = []  # Clear previous content
            st.session_state.images = []   # Clear previous images

            for goal in goals:
                st.session_state.content.append(f"Question: {goal.question}")
                st.session_state.content.append(f"Visualization: {goal.visualization}")

                sdf = SmartDataframe(temp_file_path)
                try:
                    response = sdf.chat(goal.visualization)
                    if isinstance(response, str):
                        img = Image.open(BytesIO(base64.b64decode(response.split(',')[1])))
                        st.session_state.content.append(img)
                        st.session_state.images.append(img)
                except Exception:
                    try:
                        charts = lida.visualize(summary=summary, goal=goal, textgen_config=textgen_config, library="seaborn")  
                        img_base64_string = charts[0].raster
                        img = Image.open(BytesIO(base64.b64decode(img_base64_string)))
                        st.session_state.content.append(img)
                        st.session_state.images.append(img)
                    except Exception as e:
                        st.session_state.content.append(f"Unable to generate visualization: {str(e)}")

                st.session_state.content.append(f"Rationale: {goal.rationale}")
                st.session_state.content.append("\n")  # Add a blank line between goals

        except Exception as e:
            st.error(f"An error occurred during data summarization: {str(e)}")
            st.write(f"Error details: {type(e).__name__}, {str(e)}")
        finally:
            # Clean up the temporary file
            os.unlink(temp_file_path)

    # Display the results
    for item in st.session_state.content:
        if isinstance(item, str):
            if item.startswith("Question:"):
                st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Possible question that could be asked from the data:</h1>', unsafe_allow_html=True)
            elif item.startswith("Visualization:"):
                st.markdown('<h1 style="font-size: 1rem; margin-top: 0;">Visualization:</h1>', unsafe_allow_html=True)
            st.write(item)
        elif isinstance(item, Image.Image):
            st.image(item)

    # Download options (keep this part as it was)
    if 'show_download_options' not in st.session_state:
        st.session_state.show_download_options = False

 
    if st.session_state.show_download_options:
        st.markdown("""
            <style>
            .stSelectbox {
                max-width: 200px;
            }
            </style>
            """, unsafe_allow_html=True)      
        
        file_format = st.selectbox('Select file format to download:', ['docx', 'pdf'])

        try:
            if file_format == 'docx':
                file_content = convert_to_docx(st.session_state.content, st.session_state.images)
                mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                file_name = "tabular_data_summary.docx"
            else:  # pdf
                file_content = convert_to_pdf(st.session_state.content, st.session_state.images)
                mime_type = 'application/pdf'
                file_name = "tabular_data_summary.pdf"

            st.download_button(
                label=f"Download {file_format.upper()}",
                data=file_content,
                file_name=file_name,
                mime=mime_type,
                type="primary"
            )
        except Exception as e:
            st.error(f"An error occurred while generating the download file: {str(e)}")